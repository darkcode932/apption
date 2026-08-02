import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_cahier_de_charges():
    doc = docx.Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles & Colors
    # Primary: Dark Obsidian (#0B0B0F)
    # Secondary: Emerald Green (#10B981)
    # Accent: Cyan (#06B6D4)
    # Text Dark: (#1F2937)
    
    # Title Cover Page Style
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_apption = p_title.add_run("APPTION\n")
    run_apption.font.name = "Arial"
    run_apption.font.size = Pt(36)
    run_apption.font.bold = True
    run_apption.font.color.rgb = RGBColor(16, 185, 129) # Emerald Green
    
    run_subtitle = p_title.add_run("CAHIER DES CHARGES TECHNIQUE & FONCTIONNEL\n\n")
    run_subtitle.font.name = "Arial"
    run_subtitle.font.size = Pt(20)
    run_subtitle.font.bold = True
    run_subtitle.font.color.rgb = RGBColor(11, 11, 15) # Obsidian Dark
    
    run_desc = p_title.add_run("Plateforme Mondiale d'Impact Citoyen, Pétitions Intelligentes & Amplification IA (PetBot AI)\n\n\n")
    run_desc.font.name = "Arial"
    run_desc.font.size = Pt(12)
    run_desc.font.italic = True
    run_desc.font.color.rgb = RGBColor(107, 114, 128)
    
    # Meta table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Projet :", "Plateforme Apption (Pétitions Citoyennes & IA)"),
        ("Version :", "2.4.0 (Mise à jour Continue)"),
        ("Auteur / Lead Arch :", "Équipe d'Ingénierie Apption & Agent Antigravity"),
        ("Dépôt GitHub :", "https://github.com/darkcode932/apption.git"),
        ("Dernière Mise à Jour :", "Août 2026"),
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        r0 = row.cells[0].paragraphs[0].add_run(k)
        r0.font.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = RGBColor(16, 185, 129)
        r1 = row.cells[1].paragraphs[0].add_run(v)
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(31, 41, 55)

    doc.add_page_break()
    
    # Function to add styled headings
    def add_custom_heading(text, level):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(16, 185, 129) # Emerald
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(6, 182, 212) # Cyan
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(11, 11, 15)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)

    # SECTION 1: PRESENTATION DU PROJET
    add_custom_heading("1. Présentation Générale & Vision du Projet", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Apption est une application web moderne et innovante de mobilisations citoyennes et de pétitions intelligentes. "
        "Elle ambitionne de se positionner comme le leader mondial de la démocratie participative et de l'engagement citoyen en intégrant des fonctionnalités à forte valeur ajoutée juridique, institutionnelle et médiatique.\n\n"
        "Contrairement aux plateformes classiques (Change.org, Avaaz), Apption combine l'intelligence artificielle (PetBot AI), la géolocalisation d'impact (Système GIS), "
        "des certifications juridiques et des mécanismes d'interpellation directe auprès des autorités et des médias."
    )
    
    # SECTION 2: OBJECTIFS STRATEGIQUES & METHODOLOGIE
    add_custom_heading("2. Objectifs Stratégiques & Piliers Techniques", level=1)
    bullet_points = [
        ("Design & UX Premium", "Interface sombre Obsidian (#0b0b0f), glassmorphic avec lueurs émeraude néon, typographies modernes (Outfit & Inter) et zéro astérisque markdown dans les générations d'IA."),
        ("Architecture Clean & Robustesse", "Implémentation basée sur la Clean Architecture (Domaines, Entités, Use Cases, Repositories) avec Next.js 14 (App Router) et React 18."),
        ("Données Dynamiques & Firebase", "Stockage et synchronisation temps réel avec Google Cloud Firestore (collections 'petition', 'users', 'comments', 'signatures', 'timeline')."),
        ("Sécurité & Conformité de Niveau Entreprise", "Règles Firestore strictes, validation d'intégrité, protection anti brute-force sur l'authentification, assainissement des entrées (Sanitizer XSS) et en-têtes HTTP de sécurité (Middleware Next.js)."),
        ("Fonctionnalités Premium Innovantes", "Smart Dispatch AI (Lettre recommandée électronique officielle), Amplificateur Média (Fil de Presse IA), Bouclier Juridique, Dossier Physique Relié et War Room SMS.")
    ]
    for title, desc in bullet_points:
        p = doc.add_paragraph(style='List Bullet')
        r_t = p.add_run(f"{title} : ")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(16, 185, 129)
        p.add_run(desc)

    # SECTION 3: ARCHITECTURE ET STACK TECHNIQUE
    add_custom_heading("3. Architecture & Stack Technique", level=1)
    
    tech_table = doc.add_table(rows=7, cols=2)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tech_data = [
        ("Framework Frontend / SSR", "Next.js 14.2.3 (React 18, App Router, TypeScript)"),
        ("Style & Design System", "Vanilla CSS / TailwindCSS avec variables de typographie dynamiques et Glassmorphism"),
        ("Base de Données & Auth", "Google Firebase (Cloud Firestore NoSQL, Firebase Auth avec Email/Mot de passe, Google & Facebook)"),
        ("Moteur d'Intelligence Artificielle", "Google Generative AI (Gemini 1.5 Flash - PetBot AI)"),
        ("Cartographie & GIS", "Leaflet GIS dynamique avec tuiles CartoDB Dark et rendu interactif par ville"),
        ("Sécurité & Middleware", "Next.js Middleware HTTP Security Headers, Firestore Security Rules v2, Input Sanitizer"),
        ("Internationalisation (i18n)", "Système bilingue natif Français (FR 🇫🇷) et Anglais (EN 🇬🇧)")
    ]
    for i, (k, v) in enumerate(tech_data):
        row = tech_table.rows[i]
        r0 = row.cells[0].paragraphs[0].add_run(k)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r1 = row.cells[1].paragraphs[0].add_run(v)
        r1.font.size = Pt(9.5)

    # SECTION 4: SPÉCIFICATIONS FONCTIONNELLES DÉTAILLÉES
    add_custom_heading("4. Spécifications Fonctionnelles Détaillées", level=1)
    
    add_custom_heading("4.1. Parcours Pétition & Mobilisation Citoyenne", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "• Création de Pétition en 3 étapes : Saisie du titre, description, catégorie, échelle (Ville/National/International), image, localisation géographique.\n"
        "• Objectifs & Délais Personnalisés : Fixation obligatoire d'un objectif de signatures (targetGoal) et choix optionnel d'une durée d'urgence (durationDays) avec recommandation et score d'impact calculés par PetBot AI.\n"
        "• Modale de Signature Citoyenne : Signature en 1 clic avec enregistrement du motif et de l'horodatage.\n"
        "• Débats & Témoignages : Discussion citoyenne avec système d'upvotes de soutien ('Appuyer').\n"
        "• Mur des Victoires & Célébration : Animation de célébration d'atteinte de palier et studio viral."
    )

    add_custom_heading("4.2. Assistant IA PetBot AI", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "• Chatbot Citoyen : Assistant virtuel répondeur bilingue (FR/EN) guidant l'utilisateur.\n"
        "• Formateur de Texte Propre : Génération de titres, descriptions et communiqués en texte brut sans aucun astérisque markdown (**gras**), avec mise en forme élégante (Times New Roman)."
    )

    add_custom_heading("4.3. Tableau de Bord Administration & Surveillance", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "• Dashboard Analytique Temps Réel : Suivi des 5 métriques clés (Total pétitions, drapeaux haineux IA, taux de victoires, signatures cumulées, vélocité).\n"
        "• Tour de Contrôle IA Watchtower : Modération sémantique automatique des discours haineux et propos violents avec scores de toxicité.\n"
        "• Tracker des Décideurs Institutionnels : Calcul dynamique du taux de réponse et du nombre de pétitions attribuées par autorité cible.\n"
        "• Cartographie GIS Administrative : Visualisation par ville de la densité des pétitions actives et des victoires."
    )

    # SECTION 5: FONCTIONNALITES PREMIUM INNOVANTES
    add_custom_heading("5. Gamme des Fonctionnalités Premium (Pay-per-Impact)", level=1)
    
    prem_table = doc.add_table(rows=6, cols=3)
    prem_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    prem_headers = ["Fonctionnalité", "Description & Valeur Ajoutée", "Statut Implémentation"]
    for j, h in enumerate(prem_headers):
        r = prem_table.rows[0].cells[j].paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(16, 185, 129)
        
    prem_data = [
        ("🏛️ Smart Dispatch AI (#1)", "Expédition automatisée sous forme de lettre recommandée électronique officielle certifiée aux cabinets ministériels/mairies avec accusé de réception.", "✅ Implémenté & Déployé"),
        ("⚖️ Bouclier Juridique (#2)", "Analyse de conformité réglementaire par l'IA pour zéro vice de forme et macaron 'Certifié Conforme Apption'.", "📅 Planifié"),
        ("📰 Amplificateur Média (#3)", "Rédaction par PetBot AI d'un Communiqué de Presse aux normes journalistiques et diffusion auprès de 42 rédactions partenaires.", "✅ Implémenté & Déployé"),
        ("📜 Dossier Physique Relié (#4)", "Impression d'un livre d'art officiel relié avec sceau à chaud et remise en mains propres par huissier lors des conseils municipaux.", "📅 Planifié"),
        ("📢 Campaign War Room (#5)", "Envoi d'alertes directes par SMS / WhatsApp aux signataires avec un taux d'ouverture de 98% pour mobiliser lors des votes décisifs.", "📅 Planifié")
    ]
    
    for i, row_data in enumerate(prem_data):
        row = prem_table.rows[i+1]
        for j, val in enumerate(row_data):
            r = row.cells[j].paragraphs[0].add_run(val)
            r.font.size = Pt(9)

    # SECTION 6: MATRICE DE SÉCURITÉ ET AUDIT
    add_custom_heading("6. Matrice de Sécurité & Protection du Système", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "1. Protection Firebase Rules v2 : Validation d'intégrité de createdBy, vérification que les signatures et commentaires appartiennent à l'utilisateur authentifié.\n"
        "2. Protection Anti-Brute-Force Auth : Verrouillage temporaire de 60 secondes après 5 échecs consécutifs sur les formulaires de login public et d'administration.\n"
        "3. Isolation des Clés API : process.env.GEMINI_API_KEY confiné strictement côté serveur backend.\n"
        "4. En-têtes HTTP de Sécurité : Injection par Next.js Middleware (X-Frame-Options DENY, X-Content-Type-Options nosniff, Referrer-Policy).\n"
        "5. Assainissement des Données (Input Sanitizer) : Purification systématique du texte utilisateur pour neutraliser le XSS."
    )

    # SECTION 7: HISTORIQUE ET ROADMAP DE MISE A JOUR
    add_custom_heading("7. Historique des Versions & Roadmap", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "• v1.0.0 : Lancement du socle Clean Architecture et de la création de pétitions.\n"
        "• v1.5.0 : Intégration de PetBot AI et refonte Obsidian / Dribbble du footer et de la landing page.\n"
        "• v2.0.0 : Internationalisation FR / EN complète et correctifs d'authentification Facebook & Google.\n"
        "• v2.2.0 : Intégration des objectifs de signatures (targetGoal), délais (durationDays) et décideurs cibles.\n"
        "• v2.3.0 : Correctif serveur CSS layout.css 500 et sécurisation globale (Firestore, Rate Limiter, Middleware).\n"
        "• v2.4.0 (Actuelle) : Implémentation des fonctionnalités Premium #1 (Smart Dispatch AI) et #3 (Amplificateur Média) + Génération du Cahier des Charges Word évolutif."
    )

    # Save document
    file_path = r"c:\Users\Asus\Desktop\Projects\apption\Cahier_des_Charges_Apption.docx"
    doc.save(file_path)
    print(f"Document successfully created at {file_path}")

if __name__ == "__main__":
    create_cahier_de_charges()
